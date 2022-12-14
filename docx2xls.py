#%%

import pandas as pd
from docx import Document

# Função de leituras das tabelas

def read_docx_table(document,table_num=1,nheader=1):
    table = document.tables[table_num-1]
    data = [[cell.text for cell in row.cells] for row in table.rows]
    df = pd.DataFrame(data)
    if nheader == 1:
        df = df.rename(columns = df.iloc[0]).drop(df.index[0]).reset_index(drop=True)
    elif nheader == 2:
        outside_col, inside_col = df.ilod[0], df.iloc[1]
        hier_index = pd.MultiIndex.from_tuples(list(zip(outside_col,inside_col)))
        df = pd.DataFrame(data,columns = hier_index).drop(df.index[[0,1]]).reset_index(drop=True)
    elif nheader > 2:
        print("More than two headers not currently supported")
        df = pd.DataFrame()
    return df

path = ""
document = Document(path)
table_num=1
nheader=0
df = read_docx_table(document,table_num,nheader)
df.info()

# Converter tabela do Word em dataframe 

path = ""
document = Document(path)
ntab = 41 # quantidade de tabelas
table_num=8 # tabela de início
nheader=0
df = read_docx_table(document,table_num,nheader)
    
for table_num in range(8, 41):
    df.append(read_docx_table(document,table_num,nheader))

print("Dataframe:")
# print(df)
df

print("Info do dataframe:")
print(df.info())

print("\nTamanho do dataframe:")
print(len(df))

print("\nColunas do dataframe:")
print(len(df.columns))

print("\nNúmero de linhas e colunas do dataframe:")
print(df.shape)

print("\nNúmero de elementos do dataframe:")
print(len(df))


#%%

# Listar arquivos

def get_files_list(dir_name):
    # Criar lista de arquivos e subdiretórios 
    files_list = os.listdir(dir_name)
    all_files = list()
    # Iterar sobre todas as entradas
    for entry in files_list:
        # Criar o caminho completo, a partir do nome do diretório
        full_path = os.path.join(dir_name, entry)
        # Se entrada é um diretório, resgatar a lista de arquivos neste diretório 
        if os.path.isdir(full_path):
            all_files = all_files + get_files_list(full_path)
        else:
            all_files.append(full_path)
                
    return all_files


dir_name = ""

files_list = get_files_list(dir_name)

print("\nQuantidade de arquivos totais:")
print(len(files_list))

#%%

import os, pandas as pd

df_docx_files = pd.DataFrame()

path = ""

for root, dirs, files in os.walk(path):
    for file in files:
        if file.endswith("v1.docx"):
            df_docx_files.append(file)
            print(os.path.join(root, file))

print(df_docx_files)
